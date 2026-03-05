# -*- coding: utf-8 -*-
"""Markdown to Word (.docx) conversion tool.

Converts a Markdown string (contract content) to a properly formatted
.docx file using python-docx. Heading levels map to Word styles.
"""
import logging
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from agentscope.tool import ToolResponse
from agentscope.message import TextBlock

logger = logging.getLogger(__name__)

DRAFTS_DIR = Path(os.environ.get("CONTRACT_STORAGE_DIR", Path.home() / ".copaw" / "contracts" / "drafts"))


def _cn_amount(amount: float) -> str:
    """Convert a float amount to Chinese uppercase amount string."""
    digits = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿"]

    if amount == 0:
        return "零元整"

    amount_int = int(amount)
    amount_dec = round((amount - amount_int) * 100)

    def int_to_cn(n: int) -> str:
        if n == 0:
            return ""
        parts = []
        section_list = []
        while n > 0:
            section_list.append(n % 10000)
            n //= 10000
        for i, section in enumerate(section_list):
            if section == 0:
                parts.append("零")
                continue
            s = ""
            zero = False
            for j in range(3, -1, -1):
                d = (section // (10 ** j)) % 10
                if d == 0:
                    zero = True
                else:
                    if zero:
                        s += "零"
                        zero = False
                    s += digits[d] + units[j]
            parts.append(s + big_units[i])
        return "".join(reversed(parts))

    result = int_to_cn(amount_int) + "元"
    if amount_dec > 0:
        jiao = amount_dec // 10
        fen = amount_dec % 10
        if jiao > 0:
            result += digits[jiao] + "角"
        if fen > 0:
            result += digits[fen] + "分"
    else:
        result += "整"
    return result


def _build_docx(markdown_content: str, output_path: Path) -> None:
    """Build a .docx file from Markdown content."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError(f"python-docx not installed: {e}")

    doc = Document()

    # Default page margins
    from docx.shared import Inches
    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    lines = markdown_content.strip().split("\n")
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        # Filter separator rows
        data_rows = [r for r in table_rows if not all(c.strip().startswith("-") for c in r)]
        if not data_rows:
            table_rows = []
            in_table = False
            return
        cols = max(len(r) for r in data_rows)
        tbl = doc.add_table(rows=len(data_rows), cols=cols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row[:cols]):
                cell = tbl.rows[ri].cells[ci]
                cell.text = cell_text.strip()
                if ri == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        table_rows = []
        in_table = False

    for line in lines:
        stripped = line.rstrip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            cells = [c.strip() for c in stripped[1:-1].split("|")]
            table_rows.append(cells)
            continue
        else:
            if in_table:
                flush_table()

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Horizontal rule
        elif stripped.startswith("---") and len(stripped.strip("-")) == 0:
            doc.add_paragraph("─" * 40)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered list
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        # Bold inline (simple)
        elif stripped:
            p = doc.add_paragraph()
            # Handle **bold** markers
            parts = re.split(r"\*\*(.+?)\*\*", stripped)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        else:
            # Empty line → paragraph break
            doc.add_paragraph("")

    if in_table:
        flush_table()

    doc.save(str(output_path))


async def markdown_to_word(
    markdown_content: str,
    output_filename: str = "",
    **kwargs: Any,
) -> ToolResponse:
    """Convert Markdown contract content to a Word .docx file.

    Args:
        markdown_content: Full Markdown text of the contract.
        output_filename:  Desired output filename (without path). If empty,
                          auto-generated from the first heading + timestamp.

    Returns:
        JSON string with file_path, file_size, page_count.
    """
    import json

    if not markdown_content or not markdown_content.strip():
        return ToolResponse(
            content=[TextBlock(type="text", text='{"error": "markdown_content is empty"}')]
        )

    # Auto-generate filename from first heading
    if not output_filename:
        title_match = re.search(r"^#\s+(.+)$", markdown_content, re.MULTILINE)
        title = title_match.group(1).strip() if title_match else "合同草稿"
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{title}_{ts}.docx"

    if not output_filename.endswith(".docx"):
        output_filename += ".docx"

    # Ensure output directory exists
    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = DRAFTS_DIR / output_filename

    try:
        _build_docx(markdown_content, output_path)
        file_size = output_path.stat().st_size
        result = {
            "file_path": str(output_path),
            "file_size": file_size,
            "filename": output_filename,
        }
        logger.info("markdown_to_word: created %s (%d bytes)", output_path, file_size)
        return ToolResponse(
            content=[TextBlock(type="text", text=json.dumps(result, ensure_ascii=False))]
        )
    except Exception as e:
        logger.exception("markdown_to_word failed: %s", e)
        return ToolResponse(
            content=[TextBlock(type="text", text=json.dumps({"error": str(e)}))]
        )
