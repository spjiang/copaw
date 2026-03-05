# -*- coding: utf-8 -*-
"""Local contract template search API server.

Reads real .docx/.doc template files from TEMPLATE_DIR, extracts text,
and serves a vector-search-compatible HTTP API for the template_match skill.

Run standalone:
    python local_template_api.py [port]

Environment:
    TEMPLATE_DIR  Path to the contract template root directory
                  Default: docs/合同/合同模板 (relative to project root)

API:
    POST /api/template/vector-search   → { total, templates, detected_contract_type }
    GET  /api/template/{id}/content    → { template_id, template_name, content, ... }
    GET  /api/template/list            → { total, templates }
"""
import hashlib
import json
import math
import os
import re
import sys
from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
# Default: ~/.copaw/contract_templates/
# 可通过环境变量 TEMPLATE_DIR 覆盖，例如：
#   export TEMPLATE_DIR=/path/to/your/合同模板
# 这样无论 skill 部署在哪台机器的哪个路径都能正常工作。
_DEFAULT_TEMPLATE_DIR = str(Path.home() / ".copaw" / "contract_templates")

TEMPLATE_DIR = Path(os.environ.get("TEMPLATE_DIR", _DEFAULT_TEMPLATE_DIR))

# 跳过"模版使用说明"之类的辅助文档，只索引正式合同模板
SKIP_PATTERNS = ["使用说明", "副本对应", "审批表", ".DS_Store"]


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also try tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        paragraphs.append(t)
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[extract error: {e}]"


def extract_text(path: Path) -> str:
    """Extract text from docx/doc. Falls back to filename on error."""
    ext = path.suffix.lower()
    if ext == ".docx":
        return _extract_docx(path)
    # .doc files: try python-docx2txt if available, else return filename only
    try:
        import docx2txt  # type: ignore
        return docx2txt.process(str(path)) or ""
    except Exception:
        return path.stem  # just the name as fallback


# ---------------------------------------------------------------------------
# Template index builder
# ---------------------------------------------------------------------------

def _should_skip(name: str) -> bool:
    return any(p in name for p in SKIP_PATTERNS)


def _infer_contract_type(path: Path) -> str:
    """采购 or 销售 from directory name."""
    parts = path.parts
    for p in parts:
        if "采购" in p:
            return "采购"
        if "销售" in p:
            return "销售"
    return "未知"


def _infer_sub_type(name: str) -> str:
    """Infer sub-type from file name."""
    checks = [
        ("硬件采购", ["硬件"]),
        ("软件采购", ["软件采购", "软件"]),
        ("技术开发", ["技术开发", "开发协议", "开发合同"]),
        ("技术服务", ["技术服务"]),
        ("运维服务", ["运维"]),
        ("集成服务", ["集成"]),
        ("建设工程", ["建设工程", "施工"]),
        ("咨询服务", ["咨询"]),
        ("软件销售", ["软件产品销售", "软件销售"]),
        ("综合项目", ["综合项目"]),
    ]
    for sub, kws in checks:
        if any(kw in name for kw in kws):
            return sub
    return "其他"


def _infer_is_project(name: str) -> bool:
    return "项目类" in name and "非项目" not in name


def _make_id(path: Path) -> str:
    """Stable short ID from file path hash."""
    h = hashlib.md5(str(path).encode()).hexdigest()[:8].upper()
    ct = "PUR" if "采购" in str(path) else "SAL"
    return f"TPL-{ct}-{h}"


def build_index(template_dir: Path) -> list[dict]:
    """Scan template_dir, extract text, build searchable index."""
    index = []
    if not template_dir.exists():
        print(f"[LocalTemplateAPI] WARNING: TEMPLATE_DIR not found: {template_dir}", file=sys.stderr)
        return index

    for path in sorted(template_dir.rglob("*")):
        if path.suffix.lower() not in (".docx", ".doc"):
            continue
        if _should_skip(path.name):
            continue

        contract_type = _infer_contract_type(path)
        sub_type = _infer_sub_type(path.name)
        is_project = _infer_is_project(path.name)

        # Use stem as display name (strip leading index number like "1.1")
        display_name = re.sub(r"^\d+[\.\-\+\s]+", "", path.stem).strip()
        # Strip trailing version/copy markers like " (4)"
        display_name = re.sub(r"\s*\(\d+\)\s*$", "", display_name).strip()

        text = extract_text(path)

        entry = {
            "template_id": _make_id(path),
            "template_name": display_name,
            "template_type": f"{contract_type}合同",
            "contract_type": contract_type,
            "sub_type": sub_type,
            "is_project": is_project,
            "file_path": str(path),
            "description": f"{contract_type}合同 - {sub_type}{'（项目类）' if is_project else '（非项目类）'}",
            "_text": text,          # full text for similarity (not returned to caller)
            "_text_lower": text.lower(),
        }
        index.append(entry)
        print(f"[idx] {contract_type} | {sub_type} | {display_name[:40]}")

    print(f"[LocalTemplateAPI] indexed {len(index)} templates from {template_dir}")
    return index


# Build index once at startup
_INDEX: list[dict] = []


def get_index() -> list[dict]:
    global _INDEX
    if not _INDEX:
        _INDEX = build_index(TEMPLATE_DIR)
    return _INDEX


# ---------------------------------------------------------------------------
# TF-IDF-style similarity (no external deps)
# ---------------------------------------------------------------------------

def _tokenize(text: str) -> list[str]:
    """Simple character-level 2-gram + whole-word tokenization for Chinese."""
    text = re.sub(r"\s+", "", text.lower())
    tokens = []
    for i in range(len(text) - 1):
        tokens.append(text[i:i+2])
    return tokens


def _tf(tokens: list[str]) -> dict[str, float]:
    tf: dict[str, float] = {}
    for t in tokens:
        tf[t] = tf.get(t, 0) + 1
    n = len(tokens) or 1
    return {k: v / n for k, v in tf.items()}


def _cosine(a: dict[str, float], b: dict[str, float]) -> float:
    keys = set(a) & set(b)
    if not keys:
        return 0.0
    dot = sum(a[k] * b[k] for k in keys)
    na = math.sqrt(sum(v * v for v in a.values()))
    nb = math.sqrt(sum(v * v for v in b.values()))
    return dot / (na * nb + 1e-9)


def search(query: str, contract_type: str = "", top_k: int = 5) -> dict:
    index = get_index()
    if not index:
        return {"total": 0, "templates": [], "detected_contract_type": ""}

    # ── Detect contract_type from query if not provided ───────────────────
    detected = contract_type
    if not detected:
        ql = query.lower()
        purchase_kws = ["委托方", "采购方", "甲方采购", "采购", "购买", "发包", "委托"]
        sales_kws = ["受托方", "乙方提供", "销售", "提供服务", "承接", "对外服务", "受托"]
        if any(k in ql for k in sales_kws):
            detected = "销售"
        elif any(k in ql for k in purchase_kws):
            detected = "采购"

    # ── Filter by contract_type ───────────────────────────────────────────
    pool = [t for t in index if t["contract_type"] == detected] if detected else index

    # ── Score: cosine on full text + bonus for sub_type keyword in query ──
    qtokens = _tokenize(query)
    qtf = _tf(qtokens)

    scored = []
    for tmpl in pool:
        ttf = _tf(_tokenize(tmpl["_text"]))
        sim = _cosine(qtf, ttf)

        # Sub-type keyword bonus
        sub_kws = {
            "硬件采购": ["硬件", "设备", "服务器"],
            "软件采购": ["软件", "license"],
            "技术开发": ["技术开发", "定制开发", "研发"],
            "技术服务": ["技术服务", "驻场"],
            "运维服务": ["运维", "维保", "维护"],
            "集成服务": ["集成", "系统集成"],
            "建设工程": ["建设", "施工", "工程"],
            "咨询服务": ["咨询", "顾问"],
            "软件销售": ["软件产品", "软件销售"],
            "综合项目": ["综合项目"],
        }.get(tmpl["sub_type"], [])
        ql = query.lower()
        if any(k in ql for k in sub_kws):
            sim += 0.3

        scored.append((sim, tmpl))

    scored.sort(key=lambda x: -x[0])
    results = [t for _, t in scored[:top_k]]

    # Strip internal fields before returning
    clean = [{k: v for k, v in t.items() if not k.startswith("_")} for t in results]
    return {
        "total": len(clean),
        "templates": clean,
        "detected_contract_type": detected,
    }


# ---------------------------------------------------------------------------
# HTTP server
# ---------------------------------------------------------------------------

class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args):
        pass

    def do_POST(self):
        if self.path == "/api/template/vector-search":
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length) or b"{}")
            query = body.get("query", "")
            top_k = int(body.get("top_k", 5))
            contract_type = body.get("contract_type", "")
            result = search(query, contract_type=contract_type, top_k=top_k)
            self._json(200, result)
        else:
            self._json(404, {"error": "not found"})

    def do_GET(self):
        if self.path == "/api/template/list":
            index = get_index()
            clean = [{k: v for k, v in t.items() if not k.startswith("_")} for t in index]
            self._json(200, {"total": len(clean), "templates": clean})

        elif self.path.startswith("/api/template/") and self.path.endswith("/content"):
            tid = self.path.split("/")[3]
            tmpl = next((t for t in get_index() if t["template_id"] == tid), None)
            if not tmpl:
                self._json(404, {"error": f"template {tid} not found"})
                return
            self._json(200, {
                "template_id": tid,
                "template_name": tmpl["template_name"],
                "contract_type": tmpl["contract_type"],
                "sub_type": tmpl["sub_type"],
                "file_path": tmpl["file_path"],
                "content": tmpl["_text"],
            })
        else:
            self._json(404, {"error": "not found"})

    def _json(self, code, data):
        body = json.dumps(data, ensure_ascii=False).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def run(port: int = 9000):
    # Pre-build index before accepting requests
    get_index()
    server = HTTPServer(("0.0.0.0", port), Handler)
    n = len(_INDEX)
    print(f"[LocalTemplateAPI] ready  http://0.0.0.0:{port}  ({n} templates indexed)")
    server.serve_forever()


if __name__ == "__main__":
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 9000
    run(port)
