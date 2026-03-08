#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Standalone Markdown → Word (.docx) converter for contract drafting.

No CoPaw tool registration needed. Called via execute_shell_command.

Usage:
    python word_gen.py <markdown_file> [output_dir] [exec_id] [session_id] [user_id]

    markdown_file : path to a .md file containing the contract Markdown
    output_dir    : where to save the .docx (default: ~/.copaw/contracts/drafts)
    exec_id       : optional, used as filename prefix
    session_id    : optional, used for Redis push notification
    user_id       : optional, used for Redis push notification

Output (stdout): JSON
    {
      "file_path": "/abs/path/to/合同名称.docx",
      "file_url":  "http://127.0.0.1:8088/files/contracts/drafts/合同名称.docx",
      "filename":  "合同名称.docx",
      "file_size": 12345
    }

On error, exits with code 1 and prints {"error": "..."} to stdout.
"""
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

_SHARED_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(_SHARED_DIR.parent))

from shared.push import push
from shared.redis_push import push_end, push_error, push_running, push_start

_SKILL_NAME = "合同编写智能体"


# ---------------------------------------------------------------------------
# Markdown → .docx builder
# ---------------------------------------------------------------------------

def _build_docx(markdown_text: str, output_path: Path) -> None:
    """Convert markdown to a structured Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(12)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)
        # Bullet / numbered list
        elif stripped.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+[\.\、]\s+", stripped):
            text = re.sub(r"^\d+[\.\、]\s+", "", stripped)
            doc.add_paragraph(text, style="List Number")
        # Bold lines (e.g. **甲方**：xxx)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            # Inline bold replacement: **text** → bold run
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "usage: word_gen.py <markdown_file> [output_dir] [exec_id] [session_id]"}))
        sys.exit(1)

    md_file    = Path(sys.argv[1])
    exec_id    = sys.argv[3] if len(sys.argv) > 3 else ""
    session_id = os.environ.get("COPAW_SESSION_ID") or (sys.argv[4] if len(sys.argv) > 4 else "unknown_session")
    user_id    = os.environ.get("COPAW_USER_ID") or (sys.argv[5] if len(sys.argv) > 5 else "unknown_user")
    import uuid as _uuid_mod
    run_id = str(_uuid_mod.uuid4())

    if not md_file.exists():
        err = {"error": f"markdown file not found: {md_file}"}
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data={"markdown_file": str(md_file), "exec_id": exec_id},
            error_msg=err["error"],
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps(err))
        sys.exit(1)

    markdown_text = md_file.read_text(encoding="utf-8")
    if not markdown_text.strip():
        err = {"error": "markdown file is empty"}
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data={"markdown_file": str(md_file), "exec_id": exec_id},
            error_msg=err["error"],
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps(err))
        sys.exit(1)

    # Output directory
    output_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else (
        Path(os.environ.get("CONTRACT_STORAGE_DIR", str(Path.home() / ".copaw" / "contracts")))
        / "drafts"
    )
    output_dir.mkdir(parents=True, exist_ok=True)

    # Build filename from first H1 heading
    title_match = re.search(r"^#\s+(.+)$", markdown_text, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else "合同草稿"
    title = re.sub(r'[\\/:*?"<>|]', "_", title)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    prefix = exec_id[:8] + "_" if exec_id else ""
    filename = f"{prefix}{title}_{ts}.docx"
    output_path = output_dir / filename

    # Avoid overwrite
    counter = 1
    while output_path.exists():
        output_path = output_dir / f"{prefix}{title}_{ts}_{counter}.docx"
        counter += 1

    _input = {
        "markdown_file": str(md_file),
        "exec_id": exec_id,
        "session_id": session_id,
        "title": title,
    }
    push(session_id, user_id, "📝 正在生成 Word 合同文件...", msg_type="progress")
    push_start(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        input_data=_input,
        exec_id=exec_id,
        run_id=run_id,
    )

    try:
        _build_docx(markdown_text, output_path)
    except ImportError:
        err_msg = "python-docx not installed. Run: pip install python-docx"
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err_msg,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err_msg}))
        sys.exit(1)
    except Exception as e:
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=str(e),
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

    file_size = output_path.stat().st_size

    # Construct access URL
    copaw_api = os.environ.get("COPAW_API_BASE", "http://127.0.0.1:8088")
    try:
        storage_base = Path(os.environ.get("CONTRACT_STORAGE_DIR",
                                           str(Path.home() / ".copaw" / "contracts")))
        rel = output_path.relative_to(storage_base.parent)
        file_url = f"{copaw_api}/files/{rel}"
    except ValueError:
        file_url = f"{copaw_api}/files/contracts/drafts/{filename}"

    result = {
        "file_path": str(output_path),
        "file_url":  file_url,
        "filename":  filename,
        "file_size": file_size,
    }

    push(session_id, user_id, f"✅ 合同文件已生成：{filename}", msg_type="result")
    push_running(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        render_type="llm_draft_success",
        input_data=_input,
        output_data=result,
        exec_id=exec_id,
        run_id=run_id,
    )
    push_end(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        input_data=_input,
        output_data=result,
        render_type="agent_end",
        exec_id=exec_id,
        run_id=run_id,
    )

    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
