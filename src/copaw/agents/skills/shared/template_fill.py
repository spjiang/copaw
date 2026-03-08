#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Copy a contract template and fill parameters in-place (.docx only).

Usage:
    python template_fill.py <template_file> <params_json_file> [output_dir] [exec_id] [session_id] [user_id]

Output (stdout): JSON
    {
      "file_path": "/abs/path/to/filled.docx",
      "file_url": "http://127.0.0.1:8088/files/contracts/drafts/filled.docx",
      "filename": "filled.docx",
      "file_size": 12345,
      "replaced_count": 10,
      "used_keys": ["party_a_name", "amount_cn"]
    }
"""
import json
import os
import re
import shutil
import sys
import uuid
from datetime import datetime
from pathlib import Path

_SHARED_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(_SHARED_DIR.parent))

from shared.push import push
from shared.redis_push import push_end, push_error, push_running, push_start

_SKILL_NAME = "合同模板填充"


def _load_params_map(params_file: Path) -> dict[str, str]:
    raw = params_file.read_text(encoding="utf-8").strip()
    if raw.startswith("```"):
        lines = raw.splitlines()
        raw = "\n".join(line for line in lines if not line.strip().startswith("```")).strip()
    data = json.loads(raw)

    out: dict[str, str] = {}

    def _put(k, v):
        if k is None:
            return
        key = str(k).strip()
        if not key:
            return
        if v is None:
            return
        val = str(v).strip()
        if not val:
            return
        out[key] = val

    if isinstance(data, dict) and isinstance(data.get("params"), list):
        for item in data["params"]:
            if not isinstance(item, dict):
                continue
            val = item.get("value")
            _put(item.get("field_name"), val)
            _put(item.get("label"), val)
            _put(item.get("param_id"), val)

    # Also support flat json object: {"甲方企业名称":"xxx", "party_a_name":"xxx"}
    if isinstance(data, dict):
        for k, v in data.items():
            if k == "params":
                continue
            if isinstance(v, (str, int, float, bool)):
                _put(k, v)

    return out


def _replace_by_map(text: str, params_map: dict[str, str]) -> tuple[str, int, set[str]]:
    new_text = text
    replaced_count = 0
    used_keys: set[str] = set()

    for k, v in params_map.items():
        patterns = [
            "{{" + k + "}}",
            "【" + k + "】",
            "（" + k + "）",
            "(" + k + ")",
            "[" + k + "]",
            "<" + k + ">",
            "${" + k + "}",
            "{{ " + k + " }}",
        ]
        for p in patterns:
            if p in new_text:
                cnt = new_text.count(p)
                new_text = new_text.replace(p, v)
                replaced_count += cnt
                used_keys.add(k)

    # Generic blanks: replace "【   】" by key-matched values if key words exist around.
    # Keep conservative to avoid destructive replacements.
    return new_text, replaced_count, used_keys


def _replace_paragraph(paragraph, params_map: dict[str, str]) -> tuple[int, set[str]]:
    replaced_total = 0
    used_all: set[str] = set()

    # First pass: run-level replacement preserves style for common cases.
    for run in paragraph.runs:
        old = run.text
        if not old:
            continue
        new, cnt, used = _replace_by_map(old, params_map)
        if cnt > 0:
            run.text = new
            replaced_total += cnt
            used_all |= used

    # Second pass: handle placeholders split across runs.
    joined = "".join(r.text for r in paragraph.runs)
    new_joined, cnt2, used2 = _replace_by_map(joined, params_map)
    if cnt2 > 0 and new_joined != joined:
        if paragraph.runs:
            paragraph.runs[0].text = new_joined
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.text = new_joined
        replaced_total += cnt2
        used_all |= used2

    return replaced_total, used_all


def _fill_docx(input_docx: Path, output_docx: Path, params_map: dict[str, str]) -> tuple[int, list[str]]:
    from docx import Document

    doc = Document(str(input_docx))
    replaced_count = 0
    used_keys: set[str] = set()

    for p in doc.paragraphs:
        cnt, used = _replace_paragraph(p, params_map)
        replaced_count += cnt
        used_keys |= used

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    cnt, used = _replace_paragraph(p, params_map)
                    replaced_count += cnt
                    used_keys |= used

    for section in doc.sections:
        for p in section.header.paragraphs:
            cnt, used = _replace_paragraph(p, params_map)
            replaced_count += cnt
            used_keys |= used
        for p in section.footer.paragraphs:
            cnt, used = _replace_paragraph(p, params_map)
            replaced_count += cnt
            used_keys |= used

    doc.save(str(output_docx))
    return replaced_count, sorted(used_keys)


def main():
    if len(sys.argv) < 3:
        print(json.dumps({
            "error": "usage: template_fill.py <template_file> <params_json_file> [output_dir] [exec_id] [session_id] [user_id]"
        }))
        sys.exit(1)

    template_file = Path(sys.argv[1])
    params_file = Path(sys.argv[2])
    output_dir = Path(sys.argv[3]) if len(sys.argv) > 3 else (
        Path(os.environ.get("CONTRACT_STORAGE_DIR", str(Path.home() / ".copaw" / "contracts"))) / "drafts"
    )
    exec_id = sys.argv[4] if len(sys.argv) > 4 else ""
    session_id = os.environ.get("COPAW_SESSION_ID") or (sys.argv[5] if len(sys.argv) > 5 else "unknown_session")
    user_id = os.environ.get("COPAW_USER_ID") or (sys.argv[6] if len(sys.argv) > 6 else "unknown_user")
    run_id = str(uuid.uuid4())

    _input = {
        "template_file": str(template_file),
        "params_file": str(params_file),
        "exec_id": exec_id,
    }
    push(session_id, user_id, "🧩 正在按模板填充参数并生成 Word...", msg_type="progress")
    push_start(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        input_data=_input,
        exec_id=exec_id,
        run_id=run_id,
    )

    if not template_file.exists():
        err = f"template file not found: {template_file}"
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err}, ensure_ascii=False))
        sys.exit(1)
    if template_file.suffix.lower() != ".docx":
        err = f"only .docx template is supported now: {template_file}"
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err}, ensure_ascii=False))
        sys.exit(1)
    if not params_file.exists():
        err = f"params file not found: {params_file}"
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err}, ensure_ascii=False))
        sys.exit(1)

    try:
        params_map = _load_params_map(params_file)
    except Exception as e:
        err = f"params json parse failed: {e}"
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err}, ensure_ascii=False))
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    prefix = exec_id[:8] + "_" if exec_id else ""
    filename = f"{prefix}模板填充_{template_file.stem}_{ts}.docx"
    output_path = output_dir / filename

    # Keep a copied file as requested, then fill it.
    shutil.copy2(str(template_file), str(output_path))
    push_running(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        render_type="template_render_started",
        input_data=_input,
        output_data={"output_file": str(output_path)},
        exec_id=exec_id,
        run_id=run_id,
    )

    try:
        replaced_count, used_keys = _fill_docx(output_path, output_path, params_map)
    except ImportError:
        err = "python-docx not installed. Run: pip install python-docx"
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err}, ensure_ascii=False))
        sys.exit(1)
    except Exception as e:
        err = f"fill template failed: {e}"
        push_running(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            render_type="template_render_failed",
            input_data=_input,
            output_data={"error_message": err},
            exec_id=exec_id,
            run_id=run_id,
        )
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=_SKILL_NAME,
            input_data=_input,
            error_msg=err,
            exec_id=exec_id,
            run_id=run_id,
        )
        print(json.dumps({"error": err}, ensure_ascii=False))
        sys.exit(1)

    storage_base = Path(os.environ.get("CONTRACT_STORAGE_DIR", str(Path.home() / ".copaw" / "contracts")))
    copaw_api = os.environ.get("COPAW_API_BASE", "http://127.0.0.1:8088")
    try:
        rel = output_path.relative_to(storage_base.parent)
        file_url = f"{copaw_api}/files/{rel}"
    except Exception:
        file_url = f"{copaw_api}/files/contracts/drafts/{filename}"

    result = {
        "file_path": str(output_path),
        "file_url": file_url,
        "filename": filename,
        "file_size": output_path.stat().st_size,
        "replaced_count": replaced_count,
        "used_keys": used_keys,
    }
    push(session_id, user_id, f"✅ 模板填充完成：{filename}", msg_type="result")
    push_running(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        render_type="template_render_success",
        input_data=_input,
        output_data=result,
        exec_id=exec_id,
        run_id=run_id,
    )
    push_end(
        session_id=session_id,
        user_id=user_id,
        skill_name=_SKILL_NAME,
        input_data=_input,
        output_data=result,
        render_type="agent_end",
        exec_id=exec_id,
        run_id=run_id,
    )
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
