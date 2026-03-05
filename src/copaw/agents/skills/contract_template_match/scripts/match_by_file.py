#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Match a contract template by uploaded file content.

Called when a user uploads a file in the contract drafting flow.
Extracts text from the uploaded file, then finds the most similar
templates from the real template library.

Usage:
    python match_by_file.py <session_id> <user_id> <exec_id> <file_path_or_url>

    file_path_or_url 支持两种格式：
      - 本地绝对路径：/Users/xxx/.copaw/file_store/uploads/20260305_abc.docx
      - HTTP URL    ：http://127.0.0.1:8088/api/files/download/20260305_abc.docx

    若为 HTTP URL，脚本会自动下载到临时目录后再处理。

Output (stdout): JSON
    {
      "total": N,
      "templates": [...],
      "detected_contract_type": "采购"|"销售"|"",
      "extracted_chars": N,
      "matched_file": true|false,   # true if exact file-path match found
      "local_path": "/path/to/file" # 实际使用的本地路径
    }
"""
import json
import math
import os
import re
import sys
import tempfile
import time
import uuid
import urllib.request
import urllib.parse
from datetime import datetime, timezone
from pathlib import Path


# ---------------------------------------------------------------------------
# 内联 Redis push（不依赖 shared 导入，避免路径问题导致静默失败）
# ---------------------------------------------------------------------------
def _rp(session_id, skill_name, stage, render_type,
         input_data=None, output_data=None,
         exec_id="", run_id="", runtime_ms=0):
    try:
        import redis as _r
        _cli = _r.Redis(
            host=os.environ.get("REDIS_HOST", "127.0.0.1"),
            port=int(os.environ.get("REDIS_PORT", "6379")),
            password=os.environ.get("REDIS_PASSWORD") or None,
            db=int(os.environ.get("REDIS_DB", "0")),
            socket_connect_timeout=2, socket_timeout=3, decode_responses=True,
        )
        _cli.ping()
        _payload = {
            "session_id": session_id, "exec_id": exec_id, "run_id": run_id,
            "skill_name": skill_name, "stage": stage, "render_type": render_type,
            "input": input_data or {}, "output": output_data or {},
            "timestamp": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%f")[:-3] + "Z",
            "runtime_ms": runtime_ms,
        }
        _cli.xadd(session_id, {"data": json.dumps(_payload, ensure_ascii=False)}, maxlen=500)
        print(f"[redis] → {session_id} skill={skill_name} stage={stage}", file=sys.stderr)
    except Exception as _e:
        print(f"[redis] skip: {_e}", file=sys.stderr)


def _push_bubble(session_id, user_id, msg, msg_type="progress"):
    """气泡通知（HTTP），失败静默跳过。"""
    try:
        import urllib.request as _ur, json as _j
        _body = _j.dumps({
            "session_id": session_id, "user_id": user_id,
            "text": msg, "msg_type": msg_type,
            "subscribe_id": f"{session_id}_{user_id}",
        }).encode()
        _req = _ur.Request(
            "http://127.0.0.1:8088/api/console/internal-push",
            data=_body, headers={"Content-Type": "application/json"}, method="POST",
        )
        _ur.urlopen(_req, timeout=3)
    except Exception:
        pass

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(file_path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            parts.append(t)
            return "\n".join(parts)
        except Exception as e:
            return f"[docx error: {e}]"
    elif ext == ".doc":
        # 策略1：macOS 自带 textutil（最可靠）
        try:
            import subprocess
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass
        # 策略2：antiword（Linux/brew 安装）
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass
        # 策略3：docx2txt（pip 安装，部分 .doc 格式可读）
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()
        except Exception:
            pass
        return f"[doc error: 无法解析，请安装 textutil/antiword 或转换为 .docx]"
    elif ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as f:
                return "\n".join(p.extract_text() or "" for p in f.pages)
        except Exception as e:
            return f"[pdf error: {e}]"
    elif ext in (".txt", ".md"):
        try:
            return Path(file_path).read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            return f"[read error: {e}]"
    else:
        return path.stem  # fallback to filename


# ---------------------------------------------------------------------------
# Template index (same logic as local_template_api.py)
# ---------------------------------------------------------------------------

# Default: ~/.copaw/contract_templates/
# 可通过环境变量 TEMPLATE_DIR 覆盖，与 local_template_api.py 保持一致。
_DEFAULT_TEMPLATE_DIR = str(Path.home() / ".copaw" / "contract_templates")

TEMPLATE_DIR = Path(os.environ.get("TEMPLATE_DIR", _DEFAULT_TEMPLATE_DIR))

SKIP_PATTERNS = ["使用说明", "副本对应", "审批表", ".DS_Store"]

_INDEX: list[dict] = []


def _should_skip(name: str) -> bool:
    return any(p in name for p in SKIP_PATTERNS)


def _infer_contract_type(path: Path) -> str:
    for p in path.parts:
        if "采购" in p:
            return "采购"
        if "销售" in p:
            return "销售"
    return ""


def _infer_sub_type(name: str) -> str:
    checks = [
        ("硬件采购",  ["硬件"]),
        ("软件采购",  ["软件采购", "软件"]),
        ("技术开发",  ["技术开发", "开发协议"]),
        ("技术服务",  ["技术服务"]),
        ("运维服务",  ["运维"]),
        ("集成服务",  ["集成"]),
        ("建设工程",  ["建设工程", "施工"]),
        ("咨询服务",  ["咨询"]),
        ("软件销售",  ["软件产品销售", "软件销售"]),
        ("综合项目",  ["综合项目"]),
    ]
    for sub, kws in checks:
        if any(kw in name for kw in kws):
            return sub
    return "其他"


def build_index() -> list[dict]:
    index = []
    if not TEMPLATE_DIR.exists():
        return index
    for path in sorted(TEMPLATE_DIR.rglob("*")):
        if path.suffix.lower() not in (".docx", ".doc"):
            continue
        if _should_skip(path.name):
            continue
        ct = _infer_contract_type(path)
        sub = _infer_sub_type(path.name)
        is_proj = "项目类" in path.name and "非项目" not in path.name
        display = re.sub(r"^\d+[\.\-\+\s]+", "", path.stem).strip()
        display = re.sub(r"\s*\(\d+\)\s*$", "", display).strip()
        text = extract_text(str(path))
        index.append({
            "template_id": path.stem,
            "template_name": display,
            "contract_type": ct,
            "sub_type": sub,
            "is_project": is_proj,
            "file_path": str(path),
            "_text": text,
        })
    return index


def get_index() -> list[dict]:
    global _INDEX
    if not _INDEX:
        _INDEX = build_index()
    return _INDEX


# ---------------------------------------------------------------------------
# Similarity
# ---------------------------------------------------------------------------

def _tokenize(text: str) -> list[str]:
    text = re.sub(r"\s+", "", text.lower())
    return [text[i:i+2] for i in range(len(text) - 1)]


def _tf(tokens: list[str]) -> dict[str, float]:
    tf: dict[str, float] = {}
    for t in tokens:
        tf[t] = tf.get(t, 0) + 1
    n = len(tokens) or 1
    return {k: v / n for k, v in tf.items()}


def _cosine(a: dict[str, float], b: dict[str, float]) -> float:
    keys = set(a) & set(b)
    if not keys:
        return 0.0
    dot = sum(a[k] * b[k] for k in keys)
    na = math.sqrt(sum(v * v for v in a.values()))
    nb = math.sqrt(sum(v * v for v in b.values()))
    return dot / (na * nb + 1e-9)


def _detect_contract_type(text: str) -> str:
    """
    判断合同是销售合同（软通智慧为服务提供方/受托方）还是采购合同（软通智慧为委托方/采购方）。
    优先根据"软通智慧"在合同中的角色位置判断，避免"委托"一词在销售合同中
    （"甲方委托乙方提供服务"）被误判为采购合同。
    """
    # 优先：根据软通智慧在合同中明确描述的角色
    # 软通智慧 + 乙方/提供方/受托方 → 销售合同（软通是服务提供者）
    if re.search(r"软通智慧.{0,25}(乙方|服务的提供方|提供方|受托方)", text):
        return "销售"
    if re.search(r"(乙方|服务的提供方|提供方|受托方).{0,25}软通智慧", text):
        return "销售"
    # 软通智慧 + 甲方/委托方/采购方 → 采购合同（软通是采购/委托方）
    if re.search(r"软通智慧.{0,25}(甲方|委托方|采购方)", text):
        return "采购"
    if re.search(r"(甲方|委托方|采购方).{0,25}软通智慧", text):
        return "采购"
    # 通用关键词兜底（不含软通智慧时）
    if "受托方" in text:
        return "销售"
    if "采购方" in text or "甲方采购" in text:
        return "采购"
    return ""


def find_templates(uploaded_text: str, uploaded_path: str = "", top_k: int = 5) -> dict:
    index = get_index()
    if not index:
        return {"total": 0, "templates": [], "detected_contract_type": "", "matched_file": False}

    # ── 1. 文件名匹配：剥离上传时加的 YYYYMMDD_UUID8_ 前缀，与模板库文件名精确比对 ──
    if uploaded_path:
        basename = Path(uploaded_path).name
        m = re.match(r'^\d{6,14}_[0-9a-f]{8}_(.+)$', basename)
        original_name = m.group(1) if m else basename

        for t in index:
            if Path(t["file_path"]).name == original_name:
                clean = {k: v for k, v in t.items() if not k.startswith("_")}
                return {
                    "total": 1,
                    "templates": [clean],
                    "detected_contract_type": t["contract_type"],
                    "matched_file": True,
                    "match_reason": "filename",
                }

    # ── 2. 内容相似度匹配：对全量模板库做 TF-IDF 余弦相似度，取 top_k ────────────
    detected_type = _detect_contract_type(uploaded_text)

    qtf = _tf(_tokenize(uploaded_text))
    scored = []
    for tmpl in index:
        sim = _cosine(qtf, _tf(_tokenize(tmpl["_text"])))
        scored.append((sim, tmpl))

    scored.sort(key=lambda x: -x[0])
    results = [t for _, t in scored[:top_k]]
    clean = [{k: v for k, v in t.items() if not k.startswith("_")} for t in results]

    return {
        "total": len(clean),
        "templates": clean,
        "detected_contract_type": detected_type,
        "matched_file": False,
        "match_reason": "content_similarity",
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def _is_http_url(s: str) -> bool:
    """判断字符串是否为 HTTP/HTTPS URL。"""
    try:
        parsed = urllib.parse.urlparse(s)
        return parsed.scheme in ("http", "https") and bool(parsed.netloc)
    except Exception:
        return False


def _download_to_temp(url: str) -> str:
    """
    将 HTTP URL 下载到临时文件，返回临时文件本地路径。
    临时文件名保留原始扩展名（如 .docx）。
    调用方负责在使用完后删除临时文件。
    """
    parsed = urllib.parse.urlparse(url)
    # 取 URL path 最后一段作为文件名（URL decode 处理中文）
    url_filename = urllib.parse.unquote(parsed.path.split("/")[-1]) or "upload"
    ext = os.path.splitext(url_filename)[1] or ".docx"

    tmp_dir = tempfile.mkdtemp(prefix="copaw_upload_")
    tmp_path = os.path.join(tmp_dir, url_filename)

    req = urllib.request.Request(url, headers={"User-Agent": "copaw-skill/1.0"})
    with urllib.request.urlopen(req, timeout=30) as resp:
        with open(tmp_path, "wb") as f:
            f.write(resp.read())

    if not os.path.exists(tmp_path) or os.path.getsize(tmp_path) == 0:
        raise RuntimeError(f"下载失败或文件为空：{url}")

    return tmp_path


def _pick_file_input(argv_file_input: str) -> str:
    """
    优先从 COPAW_INPUT_FILE_URLS 取附件地址，取不到时回退到入参 file_path。
    兼容:
      - JSON 数组: ["http://.../a.docx", "/path/b.docx"]
      - JSON 对象数组: [{"url":"..."}, {"path":"..."}]
      - 逗号分隔字符串: "http://.../a.docx,/path/b.docx"
      - 单值字符串: "http://.../a.docx"
    """
    raw = (os.environ.get("COPAW_INPUT_FILE_URLS") or "").strip()
    if not raw:
        return argv_file_input

    def _from_obj(item) -> str:
        if isinstance(item, str):
            return item.strip()
        if isinstance(item, dict):
            for k in ("url", "file_url", "path", "file_path"):
                v = item.get(k)
                if isinstance(v, str) and v.strip():
                    return v.strip()
        return ""

    picked = ""
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            for it in parsed:
                picked = _from_obj(it)
                if picked:
                    break
        elif isinstance(parsed, dict):
            picked = _from_obj(parsed)
        elif isinstance(parsed, str):
            picked = parsed.strip()
    except Exception:
        # 非 JSON 时按逗号分隔处理
        parts = [p.strip() for p in raw.split(",") if p.strip()]
        if parts:
            picked = parts[0]

    return picked or argv_file_input


def main():
    if len(sys.argv) < 5:
        print(json.dumps({"error": "usage: match_by_file.py session_id user_id exec_id file_path_or_url"}))
        sys.exit(1)

    session_id  = (
        os.environ.get("COPAW_SESSION_ID")
        or (sys.argv[1] if len(sys.argv) > 1 else "")
        or "unknown_session"
    )
    user_id         = os.environ.get("COPAW_USER_ID") or "unknown_user"
    exec_id         = sys.argv[3]
    file_path_input = _pick_file_input(sys.argv[4])  # 优先使用 COPAW_INPUT_FILE_URLS

    run_id      = str(uuid.uuid4())
    start_time  = time.time()
    tmp_file    = None  # 记录临时文件，用完后清理

    _skill = "合同模板匹配"
    _input = {"file_path": file_path_input, "session_id": session_id}

    _push_bubble(session_id, user_id, "🔍 正在分析上传文件，匹配合同模板...")
    _rp(session_id, _skill, "start", "progress", input_data=_input, exec_id=exec_id, run_id=run_id)

    try:
        # ------------------------------------------------------------------
        # 1. 如果是 HTTP URL，先下载到本地临时文件
        # ------------------------------------------------------------------
        if _is_http_url(file_path_input):
            _push_bubble(session_id, user_id, f"⬇️ 正在下载文件：{file_path_input}")
            try:
                tmp_file  = _download_to_temp(file_path_input)
                file_path = tmp_file
                _push_bubble(session_id, user_id, f"✅ 文件已下载到本地：{file_path}")
            except Exception as dl_err:
                _push_bubble(session_id, user_id, f"❌ 文件下载失败：{dl_err}", "error")
                err_result = {"templates": [], "total": 0,
                              "error": f"下载失败：{dl_err}",
                              "source_url": file_path_input}
                _rp(session_id, _skill, "error", "error",
                    input_data=_input, output_data={"error": str(dl_err)},
                    exec_id=exec_id, run_id=run_id,
                    runtime_ms=int((time.time() - start_time) * 1000))
                print(json.dumps(err_result))
                return
        else:
            file_path = file_path_input  # 本地路径，直接使用

        # Extract text
        _push_bubble(session_id, user_id, "📄 正在提取文件文本内容...")
        text = extract_text(file_path)
        char_count = len(text)

        if not text.strip() or text.startswith("["):
            _push_bubble(session_id, user_id, f"⚠️ 文件内容提取失败：{text}", "error")
            result = {"total": 0, "templates": [], "extracted_chars": 0, "error": text}
        else:
            _push_bubble(session_id, user_id, f"🔎 已提取 {char_count} 字，正在匹配相似模板...")
            result = find_templates(text, uploaded_path=file_path)
            result["extracted_chars"] = char_count

            total = result.get("total", 0)
            matched = result.get("matched_file", False)
            if matched:
                _push_bubble(session_id, user_id, "✅ 识别为模板库文件，已精准匹配", "result")
            elif total > 0:
                ct = result.get("detected_contract_type", "")
                _push_bubble(session_id, user_id,
                     f"✅ 找到 {total} 份相似{'（' + ct + '合同）' if ct else ''}模板", "result")
            else:
                _push_bubble(session_id, user_id, "ℹ️ 未找到相似模板，将直接基于上传文件起草", "result")

        # 在结果中附带实际使用的本地路径，方便后续步骤直接使用
        result["local_path"] = file_path
        if file_path_input != file_path:
            result["source_url"] = file_path_input

        runtime = int((time.time() - start_time) * 1000)
        _render = "template_list" if result.get("total", 0) > 0 else "progress"
        _push_bubble(session_id, user_id, f"✅ 文件模板匹配完成（{runtime}ms）", "result")
        _rp(session_id, _skill, "end", _render,
            input_data={**_input, "local_path": file_path, "extracted_chars": char_count},
            output_data={k: v for k, v in result.items() if k != "_text"},
            exec_id=exec_id, run_id=run_id, runtime_ms=runtime)

        print(json.dumps(result, ensure_ascii=False))

    except Exception as e:
        runtime = int((time.time() - start_time) * 1000)
        _push_bubble(session_id, user_id, f"❌ 模板匹配失败：{e}", "error")
        _rp(session_id, _skill, "error", "error",
            input_data=_input, output_data={"error": str(e)},
            exec_id=exec_id, run_id=run_id, runtime_ms=runtime)
        print(json.dumps({"templates": [], "total": 0, "error": str(e)}))

    finally:
        # 清理临时文件（仅当文件是从 HTTP URL 下载的临时文件时）
        if tmp_file and os.path.exists(tmp_file):
            try:
                os.remove(tmp_file)
                tmp_dir = os.path.dirname(tmp_file)
                if os.path.isdir(tmp_dir):
                    os.rmdir(tmp_dir)
            except Exception:
                pass


if __name__ == "__main__":
    main()
