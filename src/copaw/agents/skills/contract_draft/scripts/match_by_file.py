#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Template search by uploaded attachment for contract_draft."""

from __future__ import annotations

import json
import os
import sys
import tempfile
import time
import urllib.parse
import urllib.request
import uuid
from pathlib import Path

_CUR_DIR = Path(__file__).resolve().parent
_SKILLS_DIR = _CUR_DIR.parent.parent
sys.path.insert(0, str(_CUR_DIR))
sys.path.insert(0, str(_SKILLS_DIR))

from event_meta import SKILL_LABEL, SKILL_NAME, get_event_name
from push import push
from redis_push import push_end, push_error, push_running, push_start
from runtime_context import resolve_input_file_urls, resolve_session_id, resolve_user_id
from search_common import extract_keywords, infer_contract_type, search_templates

SUPPORTED_EXTS = {".docx", ".doc", ".pdf", ".txt", ".md", ".wps"}


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".docx":
        from docx import Document

        doc = Document(file_path)
        blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        blocks.append(text)
        return "\n".join(blocks)

    if ext == ".doc":
        try:
            import subprocess

            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass
        return path.stem

    if ext == ".pdf":
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")

    return path.stem


def _is_http_url(value: str) -> bool:
    parsed = urllib.parse.urlparse(value)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def _download_to_temp(url: str) -> str:
    filename = urllib.parse.unquote(urllib.parse.urlparse(url).path.split("/")[-1]) or "upload.docx"
    tmp_dir = tempfile.mkdtemp(prefix="copaw_upload_")
    target = os.path.join(tmp_dir, filename)
    request = urllib.request.Request(url, headers={"User-Agent": "copaw-skill/1.0"})
    with urllib.request.urlopen(request, timeout=30) as response:
        with open(target, "wb") as f:
            f.write(response.read())
    return target


def _pick_file_input(argv_file_input: str) -> str:
    raw = resolve_input_file_urls(argv_file_input)
    if not raw:
        return argv_file_input

    def _extract(item) -> str:
        if isinstance(item, str):
            return item.strip()
        if isinstance(item, dict):
            for key in ("url", "file_url", "path", "file_path"):
                value = item.get(key)
                if isinstance(value, str) and value.strip():
                    return value.strip()
        return ""

    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = None

    if isinstance(parsed, list):
        refs = [_extract(item) for item in parsed]
        refs = [r for r in refs if r]
        return refs[-1] if refs else argv_file_input
    if isinstance(parsed, dict):
        return _extract(parsed) or argv_file_input
    if isinstance(parsed, str) and parsed.strip():
        return parsed.strip()

    parts = [p.strip() for p in raw.split(",") if p.strip()]
    return (parts[-1] if parts else "") or argv_file_input


def main():
    if len(sys.argv) < 5:
        print(json.dumps({"error": "usage: match_by_file.py session_id user_id exec_id file_path_or_url"}))
        sys.exit(1)

    session_id = resolve_session_id(sys.argv[1] if len(sys.argv) > 1 else "")
    user_id = resolve_user_id(sys.argv[2] if len(sys.argv) > 2 else "")
    exec_id = sys.argv[3]
    file_input = _pick_file_input(sys.argv[4])
    run_id = str(uuid.uuid4())
    start_time = time.time()
    temp_file = None

    input_data = {"file_ref": file_input}
    push_start(
        session_id=session_id,
        user_id=user_id,
        skill_name=SKILL_NAME,
        skill_label=SKILL_LABEL,
        event_name=get_event_name("template_match_started"),
        input_data=input_data,
        exec_id=exec_id,
        run_id=run_id,
        render_type="template_match_started",
    )

    try:
        if _is_http_url(file_input):
            temp_file = _download_to_temp(file_input)
            file_path = temp_file
        else:
            file_path = file_input

        file_name = Path(file_path).name
        file_ext = Path(file_path).suffix.lower()
        if file_ext and file_ext not in SUPPORTED_EXTS:
            raise RuntimeError(f"不支持的文件类型: {file_ext}")

        push_running(
            session_id=session_id,
            user_id=user_id,
            skill_name=SKILL_NAME,
            skill_label=SKILL_LABEL,
            event_name=get_event_name("attachment_detected"),
            render_type="attachment_detected",
            input_data=input_data,
            output_data={
                "user_upload_contract_template_url": file_input,
                "user_upload_contract_template_file_name": file_name,
            },
            exec_id=exec_id,
            run_id=run_id,
        )

        text = extract_text(file_path)
        if not text.strip():
            raise RuntimeError("附件文本提取为空，无法进行模板匹配")

        extracted_chars = len(text)
        contract_type = infer_contract_type(text) or infer_contract_type(file_name)
        keyword_list = extract_keywords(text[:2000], filename=file_name, contract_type=contract_type)

        push_running(
            session_id=session_id,
            user_id=user_id,
            skill_name=SKILL_NAME,
            skill_label=SKILL_LABEL,
            event_name=get_event_name("user_intent_identified"),
            render_type="user_intent_identified",
            input_data={"file_name": file_name},
            output_data={
                "contract_type": contract_type,
                "keyword_list": keyword_list,
                "extracted_chars": extracted_chars,
            },
            exec_id=exec_id,
            run_id=run_id,
        )

        result = search_templates(
            query_text=" ".join(keyword_list) or text[:2000],
            keyword_list=keyword_list,
            contract_type=contract_type,
            top_k=5,
        )
        result["user_upload_contract_template_url"] = file_input
        result["user_upload_contract_template_file_name"] = file_name
        result["extracted_chars"] = extracted_chars
        result["need_user_confirm"] = result.get("total", 0) > 0
        result["need_user_intent"] = False

        runtime_ms = int((time.time() - start_time) * 1000)
        if result.get("total", 0) > 0:
            render_type = "template_candidates_found"
            push(session_id, user_id, f"✅ 已为上传附件匹配到 {result['total']} 份模板", msg_type="result")
            output_data = {
                "template_count": result["total"],
                "template_list": result["templates"],
            }
        else:
            render_type = "template_not_found"
            output_data = {"template_count": 0, "template_list": []}

        push_running(
            session_id=session_id,
            user_id=user_id,
            skill_name=SKILL_NAME,
            skill_label=SKILL_LABEL,
            event_name=get_event_name(render_type),
            render_type=render_type,
            input_data={"file_ref": file_input, "file_name": file_name, "keyword_list": keyword_list},
            output_data=output_data,
            exec_id=exec_id,
            run_id=run_id,
            runtime_ms=runtime_ms,
        )
        push_end(
            session_id=session_id,
            user_id=user_id,
            skill_name=SKILL_NAME,
            skill_label=SKILL_LABEL,
            event_name=get_event_name("template_match_finished"),
            input_data=input_data,
            output_data=result,
            render_type="template_match_finished",
            exec_id=exec_id,
            run_id=run_id,
            runtime_ms=runtime_ms,
        )
        print(json.dumps(result, ensure_ascii=False))

    except Exception as exc:
        runtime_ms = int((time.time() - start_time) * 1000)
        push(session_id, user_id, f"❌ 附件模板匹配失败：{exc}", msg_type="error")
        push_error(
            session_id=session_id,
            user_id=user_id,
            skill_name=SKILL_NAME,
            skill_label=SKILL_LABEL,
            event_name="附件模板匹配失败",
            input_data=input_data,
            error_msg=str(exc),
            exec_id=exec_id,
            run_id=run_id,
            runtime_ms=runtime_ms,
        )
        print(json.dumps({"error": str(exc), "total": 0, "templates": []}, ensure_ascii=False))
        sys.exit(1)
    finally:
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
                os.rmdir(os.path.dirname(temp_file))
            except Exception:
                pass


if __name__ == "__main__":
    main()
