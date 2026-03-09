#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Attachment parsing helpers for contract_draft."""

from __future__ import annotations

import json
import os
import subprocess
import tempfile
import urllib.parse
import urllib.request
from pathlib import Path
from typing import Any

SUPPORTED_EXTS = {
    ".docx",
    ".doc",
    ".pdf",
    ".txt",
    ".md",
    ".csv",
    ".xlsx",
}


def is_http_url(value: str) -> bool:
    parsed = urllib.parse.urlparse(value)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def download_to_temp(url: str) -> str:
    filename = urllib.parse.unquote(urllib.parse.urlparse(url).path.split("/")[-1]) or "upload.bin"
    tmp_dir = tempfile.mkdtemp(prefix="copaw_upload_")
    target = os.path.join(tmp_dir, filename)
    request = urllib.request.Request(url, headers={"User-Agent": "copaw-skill/1.0"})
    with urllib.request.urlopen(request, timeout=30) as response:
        with open(target, "wb") as file_obj:
            file_obj.write(response.read())
    return target


def _extract_ref(item: Any) -> str:
    if isinstance(item, str):
        return item.strip()
    if isinstance(item, dict):
        for key in ("url", "file_url", "path", "file_path"):
            value = item.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    return ""


def resolve_file_refs(raw: str) -> list[str]:
    if not isinstance(raw, str) or not raw.strip():
        return []

    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = None

    refs: list[str] = []
    if isinstance(parsed, list):
        for item in parsed:
            ref = _extract_ref(item)
            if ref and ref not in refs:
                refs.append(ref)
        return refs

    if isinstance(parsed, dict):
        ref = _extract_ref(parsed)
        return [ref] if ref else []

    if isinstance(parsed, str) and parsed.strip():
        return [parsed.strip()]

    for part in raw.split(","):
        ref = part.strip()
        if ref and ref not in refs:
            refs.append(ref)
    return refs


def materialize_file_ref(ref: str) -> str:
    return download_to_temp(ref) if is_http_url(ref) else ref


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".docx":
        from docx import Document

        doc = Document(file_path)
        blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    blocks.append(" | ".join(values))
        return "\n".join(blocks)

    if ext == ".doc":
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", file_path],
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
        return path.stem

    if ext == ".pdf":
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    if ext in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")

    if ext == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(file_path, read_only=True, data_only=True)
        blocks: list[str] = []
        for sheet in workbook.worksheets:
            blocks.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell not in (None, "")]
                if values:
                    blocks.append(" | ".join(values))
        workbook.close()
        return "\n".join(blocks)

    return path.stem
